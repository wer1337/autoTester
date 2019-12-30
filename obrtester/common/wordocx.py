from itertools import zip_longest
from docx.shared import Inches
from obrtester.tabs import printtest


def grouper(n, iterable, fillvalue=None):
    args = [iter(iterable)] * n
    return zip_longest(fillvalue=fillvalue, *args)


def comp_to_doc(document, image_folder, version, version2, base_path, month, day, type):
    files_list = printtest.find_files(f'{image_folder}', '.png')

    document.add_heading(f'{type}')

    for i, j in grouper(2, files_list):
        document.add_heading(f'{i[:-4]}', 1)
        document.add_paragraph(f'{version}')
        document.add_picture(f'{image_folder}\\{i}', width=Inches(6.25))
        document.add_paragraph(f'{version2}')
        document.add_picture(f'{image_folder}\\{j}', width=Inches(6.25))

    document.save(f'{base_path}\\{month}_{day}\\{type}_{version}_v_{version2}.docx')


def add_to_doc(document, img_folder1, img_folder2, version, version2, base_path, month, day, type):
    files_list = printtest.find_files(f'{img_folder1}', '.png')

    document.add_heading(f'{type}')

    for i in files_list:
        document.add_heading(f'{i[:-4]}', 1)
        document.add_paragraph(f'{version}')
        document.add_picture(f'{img_folder1}\\{i}', width=Inches(6.25))
        document.add_paragraph(f'{version2}')
        document.add_picture(f'{img_folder2}\\{i}', width=Inches(6.25))

        print(f'{img_folder1}\\{i} vs {img_folder2}\\{i}')

    document.save(f'{base_path}\\{month}_{day}\\{type}_{version}_v_{version2}_NoDiff.docx')
